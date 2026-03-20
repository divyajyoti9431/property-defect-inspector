"""
Bootstrap script: installs python-docx if needed, then creates the PTU project plan .docx.
Run with: python bootstrap_docx.py
"""
import subprocess
import sys
import os

def pip_install(package):
    print(f"Installing {package}...")
    result = subprocess.run(
        [sys.executable, "-m", "pip", "install", package, "--quiet"],
        capture_output=True, text=True
    )
    if result.returncode != 0:
        print(f"Warning: {result.stderr[:300]}")
    else:
        print(f"Installed {package} successfully")

# Install python-docx
try:
    import docx
    print("python-docx already installed")
except ImportError:
    pip_install("python-docx")
    import docx

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = r'C:\Users\divya\OBJECT DETECTION\PTU_Predictive_Tracking_Plan.docx'

# Colors
DARK_BLUE = RGBColor(0x1F, 0x4E, 0x79)
GRAY_HEADER = RGBColor(0xD9, 0xD9, 0xD9)
CODE_BG_COLOR = "F2F2F2"
GRAY_HEADER_HEX = "D9D9D9"

# ─── helper functions ───────────────────────────────────────────────────────

def set_cell_shading(cell, fill_color):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="CCCCCC"):
    """Set thin borders on all sides of a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=80, left=120, bottom=80, right=120):
    """Set cell internal margins."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)

def set_run_font(run, font_name="Arial", size_pt=11, bold=False, color=None, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # Set East Asian and complex script fonts too
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)

def set_paragraph_spacing(para, before_pt=0, after_pt=6):
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:before'), str(int(before_pt * 20)))
    spacing.set(qn('w:after'), str(int(after_pt * 20)))

def add_heading1(doc, text):
    p = doc.add_paragraph(style='Heading 1')
    set_paragraph_spacing(p, before_pt=15, after_pt=6)
    run = p.add_run(text)
    set_run_font(run, "Arial", 14, bold=True, color=DARK_BLUE)
    return p

def add_heading2(doc, text):
    p = doc.add_paragraph(style='Heading 2')
    set_paragraph_spacing(p, before_pt=10, after_pt=4)
    run = p.add_run(text)
    set_run_font(run, "Arial", 12, bold=True)
    return p

def add_normal(doc, text, bold=False, color=None, size_pt=11, align=None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=0, after_pt=6)
    run = p.add_run(text)
    set_run_font(run, "Arial", size_pt, bold=bold, color=color)
    if align:
        p.alignment = align
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    set_paragraph_spacing(p, before_pt=0, after_pt=3)
    run = p.add_run(text)
    set_run_font(run, "Arial", 11)
    return p

def add_numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    set_paragraph_spacing(p, before_pt=0, after_pt=3)
    run = p.add_run(text)
    set_run_font(run, "Arial", 11)
    return p

def add_spacer(doc):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=0, after_pt=3)
    return p

def add_code_block(doc, lines):
    """Add a code block as a single-cell table with gray background."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Set table width to full content width
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '9360')
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    cell = table.cell(0, 0)
    # Set cell width
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), '9360')
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

    set_cell_shading(cell, CODE_BG_COLOR)
    set_cell_borders(cell)
    set_cell_margins(cell, top=80, left=160, bottom=80, right=160)

    # Clear default paragraph and add code lines
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)

    for line in lines:
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        pPr.append(spacing)
        p.append(pPr)

        if line:  # non-empty line
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Courier New')
            rFonts.set(qn('w:hAnsi'), 'Courier New')
            rFonts.set(qn('w:cs'), 'Courier New')
            rPr.append(rFonts)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '18')
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), '18')
            rPr.append(szCs)
            r.append(rPr)
            t = OxmlElement('w:t')
            if line.startswith(' ') or line.endswith(' ') or '  ' in line:
                t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = line
            r.append(t)
            p.append(r)
        else:
            # Empty line - add a run with a space
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Courier New')
            rFonts.set(qn('w:hAnsi'), 'Courier New')
            rFonts.set(qn('w:cs'), 'Courier New')
            rPr.append(rFonts)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '18')
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), '18')
            rPr.append(szCs)
            r.append(rPr)
            t = OxmlElement('w:t')
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = ' '
            r.append(t)
            p.append(r)

        tc.append(p)

    return table

def add_data_table(doc, headers, rows, col_widths_dxa):
    """Add a professional data table with gray header row."""
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set table width
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    total_w = sum(col_widths_dxa)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(total_w))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    # Set grid columns
    tblGrid = OxmlElement('w:tblGrid')
    for w in col_widths_dxa:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(w))
        tblGrid.append(gridCol)
    tbl.insert(1, tblGrid)

    def fill_cell(cell, text, width_dxa, is_header=False):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(width_dxa))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)
        set_cell_borders(cell)
        set_cell_margins(cell)
        if is_header:
            set_cell_shading(cell, GRAY_HEADER_HEX)

        # Set text
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, 0, 0)
        run = p.add_run(text)
        set_run_font(run, "Arial", 10, bold=is_header)

    # Header row
    hdr_row = table.rows[0]
    for i, (header, width) in enumerate(zip(headers, col_widths_dxa)):
        fill_cell(hdr_row.cells[i], header, width, is_header=True)

    # Data rows
    for ri, row_data in enumerate(rows):
        data_row = table.rows[ri + 1]
        for ci, (cell_text, width) in enumerate(zip(row_data, col_widths_dxa)):
            fill_cell(data_row.cells[ci], cell_text, width, is_header=False)

    return table

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)
    set_paragraph_spacing(p, 0, 0)
    return p

def set_page_size_margins(doc):
    """Set US Letter page size with 1-inch margins."""
    section = doc.sections[0]
    section.page_width = Emu(12240 * 914400 // 1440)   # 12240 DXA in EMU
    section.page_height = Emu(15840 * 914400 // 1440)  # 15840 DXA in EMU
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

def add_footer_page_numbers(doc):
    """Add page number footer."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    # Clear existing paragraphs
    for para in footer.paragraphs:
        p = para._p
        p.getparent().remove(p)

    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    p.append(pPr)

    def add_text_run(parent, text):
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Arial')
        rFonts.set(qn('w:hAnsi'), 'Arial')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '18')
        rPr.append(sz)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = text
        r.append(t)
        parent.append(r)

    def add_field_run(parent, field_name):
        # begin
        r1 = OxmlElement('w:r')
        rPr1 = OxmlElement('w:rPr')
        rFonts1 = OxmlElement('w:rFonts')
        rFonts1.set(qn('w:ascii'), 'Arial')
        rFonts1.set(qn('w:hAnsi'), 'Arial')
        rPr1.append(rFonts1)
        sz1 = OxmlElement('w:sz')
        sz1.set(qn('w:val'), '18')
        rPr1.append(sz1)
        r1.append(rPr1)
        fc1 = OxmlElement('w:fldChar')
        fc1.set(qn('w:fldCharType'), 'begin')
        r1.append(fc1)
        parent.append(r1)
        # instrText
        r2 = OxmlElement('w:r')
        rPr2 = OxmlElement('w:rPr')
        rFonts2 = OxmlElement('w:rFonts')
        rFonts2.set(qn('w:ascii'), 'Arial')
        rFonts2.set(qn('w:hAnsi'), 'Arial')
        rPr2.append(rFonts2)
        sz2 = OxmlElement('w:sz')
        sz2.set(qn('w:val'), '18')
        rPr2.append(sz2)
        r2.append(rPr2)
        instr = OxmlElement('w:instrText')
        instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        instr.text = f' {field_name} '
        r2.append(instr)
        parent.append(r2)
        # end
        r3 = OxmlElement('w:r')
        rPr3 = OxmlElement('w:rPr')
        rFonts3 = OxmlElement('w:rFonts')
        rFonts3.set(qn('w:ascii'), 'Arial')
        rFonts3.set(qn('w:hAnsi'), 'Arial')
        rPr3.append(rFonts3)
        sz3 = OxmlElement('w:sz')
        sz3.set(qn('w:val'), '18')
        rPr3.append(sz3)
        r3.append(rPr3)
        fc3 = OxmlElement('w:fldChar')
        fc3.set(qn('w:fldCharType'), 'end')
        r3.append(fc3)
        parent.append(r3)

    add_text_run(p, 'Page ')
    add_field_run(p, 'PAGE')
    add_text_run(p, ' of ')
    add_field_run(p, 'NUMPAGES')

    footer._element.body.append(p)

# ─── build document ─────────────────────────────────────────────────────────

def build_document():
    doc = Document()
    set_page_size_margins(doc)
    add_footer_page_numbers(doc)

    # ── Title ──
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=24, after_pt=6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PTU-Based Laser Pointing System")
    set_run_font(run, "Arial", 26, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=0, after_pt=24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("From Reactive to Predictive Tracking \u2014 Complete Project Plan & Pipeline")
    set_run_font(run, "Arial", 13, color=RGBColor(0x55, 0x55, 0x55))

    # ── SECTION 1 ──
    add_heading1(doc, "SECTION 1: PROJECT OVERVIEW")

    add_heading2(doc, "1.1 Objective")
    add_normal(doc, "Replace the existing reactive \u201creduce-the-distance\u201d PTU tracking logic with a predictive, Kalman Filter-based algorithm to achieve \u226580% laser-on-drone time during active movement.")
    add_spacer(doc)

    add_heading2(doc, "1.2 Existing Infrastructure (Do Not Modify)")
    for item in ["Working UI (User Interface)",
                 "Camera integration with 36x optical zoom",
                 "PTU (Pan-Tilt Unit) hardware control",
                 "YOLOv8-based drone detection pipeline",
                 "Basic reactive tracking loop"]:
        add_bullet(doc, item)
    add_spacer(doc)

    add_heading2(doc, "1.3 Success Criteria")
    add_data_table(doc,
        ["Metric", "Target"],
        [["Drone speed handled", "Up to 1 m/s"],
         ["Drone distance", "Up to 50 meters"],
         ["Laser on-target rate", "\u226580% during movement"],
         ["System latency", "Minimized to meet 80% target"]],
        [4680, 4680])
    add_spacer(doc)

    # ── SECTION 2 ──
    add_page_break(doc)
    add_heading1(doc, "SECTION 2: CURRENT SYSTEM ANALYSIS (Phase 1)")

    add_heading2(doc, "2.1 Reactive Tracking \u2014 The Problem")
    add_normal(doc, "Current logic (pseudo-code):")
    add_spacer(doc)
    add_code_block(doc, [
        "loop:",
        "    bbox = yolo_detect(frame)",
        "    dx = bbox.center_x - crosshair_x",
        "    dy = bbox.center_y - crosshair_y",
        "    ptu.move(step * sign(dx), step * sign(dy))"
    ])
    add_spacer(doc)
    add_normal(doc, "Problems:")
    for item in ["No velocity estimation \u2014 the PTU always reacts to where the drone WAS",
                 "Constant step size causes oscillation near target",
                 "No prediction \u2014 laser lags behind moving drone",
                 "No Kalman smoothing \u2014 noisy detections cause jitter"]:
        add_bullet(doc, item)
    add_spacer(doc)

    add_heading2(doc, "2.2 Latency Sources to Measure")
    for item in ["YOLO inference time (typically 20\u201350 ms on GPU, 80\u2013150 ms on CPU)",
                 "Frame capture/buffer delay",
                 "PTU command round-trip time (serial/TCP)",
                 "Python GIL contention in main loop"]:
        add_bullet(doc, item)
    add_spacer(doc)

    add_heading2(doc, "2.3 Analysis Tasks")
    for item in ["Instrument the existing loop with timestamps at each stage",
                 "Log: frame_time, detection_time, ptu_command_time, ptu_ack_time",
                 "Calculate total lag = (ptu_ack_time) - (drone_actual_position_time)",
                 "Identify the dominant bottleneck"]:
        add_numbered(doc, item)
    add_spacer(doc)

    # ── SECTION 3 ──
    add_page_break(doc)
    add_heading1(doc, "SECTION 3: IMPLEMENTATION PLAN (Phase 2)")

    add_heading2(doc, "3.1 Architecture Overview")
    add_normal(doc, "The new pipeline:")
    add_spacer(doc)
    add_code_block(doc, ["Camera Frame \u2192 YOLOv8 Detection \u2192 Kalman Filter \u2192 Predictive Lookahead \u2192 Coordinate Transform \u2192 PTU Command"])
    add_spacer(doc)

    add_heading2(doc, "3.2 Component 1 \u2014 Kalman Filter (State Estimator)")
    add_normal(doc, "State vector: [x, y, vx, vy] (pixel position + velocity)")
    add_spacer(doc)
    add_normal(doc, "The Kalman Filter will:")
    for item in ["Estimate the drone\u2019s true position from noisy YOLO bounding boxes",
                 "Predict the drone\u2019s state one step ahead (even when detection fails)",
                 "Smooth out jitter from YOLO fluctuations"]:
        add_bullet(doc, item)
    add_spacer(doc)
    add_normal(doc, "Key equations:")
    add_bullet(doc, "Predict: x\u0302 = F\u00b7x + B\u00b7u (state transition)")
    add_bullet(doc, "Update: x\u0302 = x\u0302 + K\u00b7(z - H\u00b7x\u0302) (measurement correction)")
    add_spacer(doc)
    add_normal(doc, "Recommended library: filterpy (pip install filterpy)")
    add_spacer(doc)

    add_heading2(doc, "3.3 Component 2 \u2014 Predictive Lookahead")
    add_normal(doc, "After the Kalman Filter estimates [x, y, vx, vy], aim at the FUTURE position:")
    add_spacer(doc)
    add_code_block(doc, [
        "lookahead_time = system_lag_ms / 1000.0  # measured in Phase 1",
        "predicted_x = kalman_x + vx * lookahead_time",
        "predicted_y = kalman_y + vy * lookahead_time"
    ])
    add_spacer(doc)
    add_normal(doc, "The lookahead_time compensates exactly for the measured system lag.")
    add_spacer(doc)

    add_heading2(doc, "3.4 Component 3 \u2014 Coordinate Transformation")
    add_normal(doc, "Convert pixel coordinates to PTU pan/tilt angles:")
    add_spacer(doc)
    add_code_block(doc, [
        "pan_angle  = (predicted_x - frame_center_x) * (H_FOV / frame_width)",
        "tilt_angle = (predicted_y - frame_center_y) * (V_FOV / frame_height)"
    ])
    add_spacer(doc)
    add_normal(doc, "Where:")
    for item in ["H_FOV = horizontal field of view (degrees) at current zoom level",
                 "V_FOV = vertical field of view (degrees) at current zoom level",
                 "These must be calibrated per zoom level (36x zoom \u2248 ~0.5\u00b0 \u00d7 0.3\u00b0 FOV)"]:
        add_bullet(doc, item)
    add_spacer(doc)

    add_heading2(doc, "3.5 Component 4 \u2014 Adaptive Step / PID Controller (Optional Enhancement)")
    add_normal(doc, "Replace fixed step moves with a PID controller:")
    add_spacer(doc)
    add_code_block(doc, [
        "error_pan  = target_pan  - current_pan",
        "error_tilt = target_tilt - current_tilt",
        "ptu_pan_cmd  = Kp*error_pan  + Ki*integral_pan  + Kd*derivative_pan",
        "ptu_tilt_cmd = Kp*error_tilt + Ki*integral_tilt + Kd*derivative_tilt"
    ])
    add_spacer(doc)
    add_normal(doc, "Benefits:")
    for item in ["Proportional: faster response to large errors",
                 "Integral: eliminates steady-state offset",
                 "Derivative: damping to prevent overshoot"]:
        add_bullet(doc, item)
    add_spacer(doc)

    # ── SECTION 4 ──
    add_page_break(doc)
    add_heading1(doc, "SECTION 4: COMPLETE CODE PIPELINE")

    add_heading2(doc, "4.1 File Structure")
    add_code_block(doc, [
        "project/",
        "\u251c\u2500\u2500 tracker/",
        "\u2502   \u251c\u2500\u2500 __init__.py",
        "\u2502   \u251c\u2500\u2500 kalman_tracker.py      \u2190 Kalman Filter wrapper",
        "\u2502   \u251c\u2500\u2500 predictor.py           \u2190 Lookahead logic",
        "\u2502   \u251c\u2500\u2500 coord_transform.py     \u2190 Pixel \u2192 PTU angle",
        "\u2502   \u2514\u2500\u2500 pid_controller.py      \u2190 Optional PID",
        "\u251c\u2500\u2500 main_tracking_loop.py      \u2190 Replace existing loop",
        "\u2514\u2500\u2500 config.py                  \u2190 All tunable parameters"
    ])
    add_spacer(doc)

    add_heading2(doc, "4.2 kalman_tracker.py")
    add_code_block(doc, [
        "import numpy as np",
        "from filterpy.kalman import KalmanFilter",
        "",
        "class DroneKalmanTracker:",
        "    def __init__(self, dt=0.033):  # dt = 1/30 fps",
        "        self.kf = KalmanFilter(dim_x=4, dim_z=2)",
        "        self.dt = dt",
        "        # State: [x, y, vx, vy]",
        "        self.kf.F = np.array([[1,0,dt,0],",
        "                               [0,1,0,dt],",
        "                               [0,0,1, 0],",
        "                               [0,0,0, 1]])",
        "        # Measurement: [x, y]",
        "        self.kf.H = np.array([[1,0,0,0],",
        "                               [0,1,0,0]])",
        "        self.kf.R *= 10    # Measurement noise",
        "        self.kf.Q *= 0.1   # Process noise",
        "        self.initialized = False",
        "",
        "    def update(self, cx, cy):",
        "        if not self.initialized:",
        "            self.kf.x = np.array([cx, cy, 0, 0])",
        "            self.initialized = True",
        "        else:",
        "            self.kf.predict()",
        "            self.kf.update(np.array([cx, cy]))",
        "        return self.kf.x  # [x, y, vx, vy]",
        "",
        "    def predict_only(self):",
        '        """Call when detection fails -- extrapolate position"""',
        "        self.kf.predict()",
        "        return self.kf.x"
    ])
    add_spacer(doc)

    add_heading2(doc, "4.3 predictor.py")
    add_code_block(doc, [
        "class PredictiveLookahead:",
        "    def __init__(self, lag_seconds=0.1):",
        "        self.lag = lag_seconds  # Measured system lag",
        "",
        "    def get_target(self, state):",
        "        x, y, vx, vy = state",
        "        target_x = x + vx * self.lag",
        "        target_y = y + vy * self.lag",
        "        return target_x, target_y"
    ])
    add_spacer(doc)

    add_heading2(doc, "4.4 coord_transform.py")
    add_code_block(doc, [
        "class CoordTransform:",
        "    def __init__(self, frame_w, frame_h, h_fov_deg, v_fov_deg):",
        "        self.fw = frame_w",
        "        self.fh = frame_h",
        "        self.h_fov = h_fov_deg",
        "        self.v_fov = v_fov_deg",
        "",
        "    def pixel_to_angle(self, px, py):",
        "        cx, cy = self.fw / 2, self.fh / 2",
        "        pan   = (px - cx) / self.fw * self.h_fov",
        "        tilt  = (py - cy) / self.fh * self.v_fov",
        "        return pan, tilt"
    ])
    add_spacer(doc)

    add_heading2(doc, "4.5 main_tracking_loop.py (New Logic)")
    add_code_block(doc, [
        "tracker = DroneKalmanTracker(dt=1/fps)",
        "predictor = PredictiveLookahead(lag_seconds=measured_lag)",
        "transform = CoordTransform(1920, 1080, h_fov=0.5, v_fov=0.3)",
        "",
        "while True:",
        "    frame = camera.get_frame()",
        "    detections = yolo.detect(frame)",
        "",
        "    if detections:",
        "        cx, cy = get_bbox_center(detections[0])",
        "        state = tracker.update(cx, cy)",
        "    else:",
        "        state = tracker.predict_only()  # Coasting",
        "",
        "    if tracker.initialized:",
        "        target_x, target_y = predictor.get_target(state)",
        "        pan, tilt = transform.pixel_to_angle(target_x, target_y)",
        "        ptu.move_absolute(pan, tilt)"
    ])
    add_spacer(doc)

    # ── SECTION 5 ──
    add_page_break(doc)
    add_heading1(doc, "SECTION 5: OPTIMIZATION (Phase 3)")

    add_heading2(doc, "5.1 Python Performance Optimizations")
    for item in ["Run YOLO inference in a separate thread/process to avoid blocking the PTU command loop",
                 "Use a double-buffer: YOLO thread writes detection; main thread reads + sends PTU commands",
                 "Use numpy for all matrix operations (already used by filterpy)",
                 "Set Python process priority to high (os.nice(-10) on Linux)",
                 "Use asyncio or threading.Thread for non-blocking PTU communication"]:
        add_bullet(doc, item)
    add_spacer(doc)

    add_heading2(doc, "5.2 Threading Architecture")
    add_code_block(doc, [
        "Thread 1 (YOLO):    camera \u2192 yolo_detect \u2192 write to shared_detection (lock-free ring buffer)",
        "Thread 2 (Control): read shared_detection \u2192 kalman_update \u2192 predict \u2192 ptu_command"
    ])
    add_spacer(doc)

    add_heading2(doc, "5.3 Kalman Tuning Parameters")
    add_data_table(doc,
        ["Parameter", "Effect", "Tune When"],
        [["R (measurement noise)", "Higher = trust model more", "YOLO detections are noisy/jittery"],
         ["Q (process noise)", "Higher = trust measurements more", "Drone changes direction suddenly"],
         ["dt (time step)", "Must match actual frame rate", "FPS changes"],
         ["lag_seconds", "Lookahead compensation", "Measured system lag changes"]],
        [2500, 3430, 3430])
    add_spacer(doc)

    # ── SECTION 6 ──
    add_page_break(doc)
    add_heading1(doc, "SECTION 6: TESTING & VALIDATION")

    add_heading2(doc, "6.1 Test Protocol")
    for item in ["Record ground truth video with known drone trajectory (ruler-measured movement at 1 m/s)",
                 "Run new tracking code on recorded video",
                 "For each frame: log (laser_position, drone_position, on_target = distance < threshold)",
                 "Calculate success_rate = frames_on_target / total_frames"]:
        add_numbered(doc, item)
    add_spacer(doc)

    add_heading2(doc, "6.2 Success Metrics")
    add_data_table(doc,
        ["Test", "Pass Condition"],
        [["Static drone", "100% on-target within 2 seconds"],
         ["Slow linear (0.3 m/s)", "\u226590% on-target"],
         ["Fast linear (1 m/s)", "\u226580% on-target"],
         ["Direction change", "\u22640.5s recovery time"],
         ["Detection dropout (0.5s)", "Resumes tracking after re-detection"]],
        [4680, 4680])
    add_spacer(doc)

    add_heading2(doc, "6.3 Validation Script (pseudo-code)")
    add_code_block(doc, [
        "results = []",
        "for frame, gt_pos in zip(video_frames, ground_truth):",
        "    detections = yolo.detect(frame)",
        "    # ... run tracker ...",
        "    laser_pos = ptu.get_current_position()",
        "    on_target = distance(laser_pos, gt_pos) < THRESHOLD_PIXELS",
        "    results.append(on_target)",
        "",
        "success_rate = sum(results) / len(results)",
        'print(f"On-target rate: {success_rate:.1%}")',
        'assert success_rate >= 0.80, "Target not met -- tune Kalman parameters"'
    ])
    add_spacer(doc)

    # ── SECTION 7 ──
    add_page_break(doc)
    add_heading1(doc, "SECTION 7: IMPLEMENTATION TIMELINE")

    add_data_table(doc,
        ["Week", "Phase", "Tasks"],
        [["Week 1", "Phase 1: Analysis", "Instrument loop, measure latency, document interfaces"],
         ["Week 2", "Phase 2: Kalman", "Implement KalmanTracker, unit test with synthetic data"],
         ["Week 3", "Phase 2: Integration", "Integrate predictor + coordinate transform, connect to PTU"],
         ["Week 4", "Phase 3: Testing", "Run validation tests, tune parameters, achieve \u226580%"],
         ["Week 5", "Phase 3: Optimization", "Threading, performance profiling, final validation"]],
        [1560, 2600, 5200])
    add_spacer(doc)

    # ── SECTION 8 ──
    add_page_break(doc)
    add_heading1(doc, "SECTION 8: DEPENDENCIES & SETUP")

    add_heading2(doc, "8.1 Python Packages")
    add_code_block(doc, ["pip install filterpy numpy opencv-python ultralytics"])
    add_spacer(doc)

    add_heading2(doc, "8.2 Key Configuration (config.py)")
    add_code_block(doc, [
        "# Camera",
        "FRAME_WIDTH  = 1920",
        "FRAME_HEIGHT = 1080",
        "FPS = 30",
        "",
        "# Zoom-calibrated FOV (measure for each zoom level)",
        "H_FOV_DEG = 0.5   # 36x zoom horizontal FOV",
        "V_FOV_DEG = 0.3   # 36x zoom vertical FOV",
        "",
        "# Kalman tuning",
        "KALMAN_R = 10.0",
        "KALMAN_Q = 0.1",
        "",
        "# System lag (measure in Phase 1)",
        "SYSTEM_LAG_SECONDS = 0.10",
        "",
        "# On-target threshold",
        "ON_TARGET_PIXEL_RADIUS = 20"
    ])
    add_spacer(doc)

    # ── SECTION 9 ──
    add_page_break(doc)
    add_heading1(doc, "SECTION 9: RISKS & MITIGATIONS")

    add_data_table(doc,
        ["Risk", "Mitigation"],
        [["FOV calibration error at 36x zoom", "Calibrate FOV per zoom level using known target distance"],
         ["Kalman diverges on fast maneuvers", "Increase Q (process noise) to track sudden acceleration"],
         ["YOLO detection gaps > 0.5s", "Implement coast mode with velocity decay"],
         ["PTU command latency spike", "Use non-blocking PTU commands; measure and compensate"],
         ["Python GIL limits throughput", "Separate YOLO to subprocess with multiprocessing.Queue"]],
        [4680, 4680])
    add_spacer(doc)
    add_spacer(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, 0, 0)
    run = p.add_run("END OF DOCUMENT")
    set_run_font(run, "Arial", 11, bold=True, color=DARK_BLUE)

    return doc

# ─── main ───────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    print("Building document...")
    doc = build_document()
    print(f"Saving to {OUTPUT_PATH}...")
    doc.save(OUTPUT_PATH)
    size = os.path.getsize(OUTPUT_PATH)
    print(f"SUCCESS: {OUTPUT_PATH} ({size:,} bytes)")
